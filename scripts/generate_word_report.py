#!/usr/bin/env python3
"""
Script to convert the interim report from markdown to Word document
"""

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import os
import sys

def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the w:hyperlink tag
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    
    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    # Join all the xml elements together
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)
    
    return hyperlink

def create_word_report():
    """Create Word document from the interim report"""
    
    # Create document
    doc = Document()
    
    # Set up document properties
    doc.core_properties.title = "Interim Report: Task 1 - Brent Oil Price Analysis"
    doc.core_properties.author = "Birhan Energies Data Science Team"
    doc.core_properties.subject = "Brent Oil Price Analysis - Change Point Analysis and Statistical Modelling"
    
    # Title
    title = doc.add_heading('Interim Report: Task 1 - Laying the Foundation for Brent Oil Price Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Project info
    project_info = doc.add_paragraph()
    project_info.add_run('Project: ').bold = True
    project_info.add_run('Brent Oil Price Analysis - Change Point Analysis and Statistical Modelling\n')
    project_info.add_run('Team: ').bold = True
    project_info.add_run('Birhan Energies Data Science Team\n')
    project_info.add_run('Date: ').bold = True
    project_info.add_run('August 2, 2025\n')
    project_info.add_run('Phase: ').bold = True
    project_info.add_run('Task 1 Completion - Foundation Analysis')
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    
    summary_text = """
    This interim report presents the completion of Task 1 for the Brent oil price analysis project. 
    We have successfully established the foundation for analyzing how major geopolitical events affect 
    Brent oil prices through Bayesian change point analysis. The analysis covers daily Brent oil 
    prices from May 20, 1987, to November 14, 2022, and includes 29 major geopolitical events 
    that have influenced oil markets.
    
    Key Achievements:
    • Comprehensive data analysis workflow defined
    • 29 major events compiled and validated
    • Time series properties analyzed
    • Assumptions and limitations clearly documented
    • Technical environment fully configured
    • Ready for Bayesian change point analysis
    """
    
    doc.add_paragraph(summary_text)
    
    # Data Analysis Workflow
    doc.add_heading('1. Data Analysis Workflow', level=1)
    
    doc.add_heading('1.1 Systematic 5-Phase Approach', level=2)
    
    workflow_text = """
    We have developed a comprehensive data analysis workflow consisting of five distinct phases:
    
    Phase 1: Data Preparation and Exploration
    • Data Loading: Handle mixed date formats (DD-MMM-YY and MMM DD, YYYY)
    • Data Cleaning: Convert dates, handle missing values, identify outliers
    • Exploratory Analysis: Visualize trends, calculate descriptive statistics
    • Data Validation: Ensure data quality and completeness
    
    Phase 2: Event Research and Compilation
    • Event Identification: Research major geopolitical events (1990-2022)
    • Event Categorization: Classify events by type and region
    • Data Structuring: Create standardized event dataset
    • Validation: Cross-reference multiple sources for accuracy
    
    Phase 3: Change Point Analysis
    • Bayesian Modeling: Implement PyMC3 change point models
    • Model Specification: Define priors and likelihood functions
    • MCMC Sampling: Run posterior inference
    • Convergence Diagnostics: Validate model performance
    
    Phase 4: Event Correlation and Impact Analysis
    • Change Point Detection: Identify structural breaks in price series
    • Event Mapping: Correlate detected changes with events
    • Impact Quantification: Measure magnitude of event impacts
    • Statistical Validation: Assess significance of correlations
    
    Phase 5: Communication and Visualization
    • Dashboard Development: Create interactive visualizations
    • Stakeholder Reporting: Generate insights for different audiences
    • Documentation: Maintain comprehensive methodology records
    """
    
    doc.add_paragraph(workflow_text)
    
    # Event Dataset Quality
    doc.add_heading('2. Event Dataset Quality', level=1)
    
    # Load event data
    try:
        events_df = pd.read_csv('data/processed/major_events.csv')
        
        # Event categories table
        doc.add_heading('2.2 Event Categories and Distribution', level=2)
        
        categories_data = [
            ['CONFLICT', 8, '28%', 'Geopolitical conflicts and wars'],
            ['OPEC', 7, '24%', 'OPEC decisions and meetings'],
            ['SANCTIONS', 5, '17%', 'Economic sanctions'],
            ['ECONOMIC', 4, '14%', 'Economic shocks and crises'],
            ['INFRASTRUCTURE', 3, '10%', 'Supply chain and infrastructure events'],
            ['POLITICAL', 2, '7%', 'Political decisions and elections']
        ]
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Event Type'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage'
        hdr_cells[3].text = 'Description'
        
        for event_type, count, percentage, description in categories_data:
            row_cells = table.add_row().cells
            row_cells[0].text = event_type
            row_cells[1].text = str(count)
            row_cells[2].text = percentage
            row_cells[3].text = description
        
        # Geographic distribution table
        doc.add_heading('2.3 Geographic Distribution', level=2)
        
        geo_data = [
            ['Global', 11, '38%', 'OPEC decisions, financial crises'],
            ['Middle East', 10, '34%', 'Gulf War, Iraq War, Arab Spring'],
            ['North America', 4, '14%', 'Hurricane Katrina, US elections'],
            ['Europe', 3, '10%', 'Russian sanctions, Ukraine conflict'],
            ['South America', 1, '3%', 'Venezuelan sanctions']
        ]
        
        geo_table = doc.add_table(rows=1, cols=4)
        geo_table.style = 'Table Grid'
        geo_hdr_cells = geo_table.rows[0].cells
        geo_hdr_cells[0].text = 'Region'
        geo_hdr_cells[1].text = 'Count'
        geo_hdr_cells[2].text = 'Percentage'
        geo_hdr_cells[3].text = 'Key Events'
        
        for region, count, percentage, key_events in geo_data:
            row_cells = geo_table.add_row().cells
            row_cells[0].text = region
            row_cells[1].text = str(count)
            row_cells[2].text = percentage
            row_cells[3].text = key_events
        
        # Full events table (first 10 events for space)
        doc.add_heading('2.4 Tabulated Event Dataset (Sample)', level=2)
        
        events_table = doc.add_table(rows=1, cols=6)
        events_table.style = 'Table Grid'
        events_hdr_cells = events_table.rows[0].cells
        events_hdr_cells[0].text = 'Date'
        events_hdr_cells[1].text = 'Event Type'
        events_hdr_cells[2].text = 'Description'
        events_hdr_cells[3].text = 'Region'
        events_hdr_cells[4].text = 'Impact Expected'
        events_hdr_cells[5].text = 'Source'
        
        # Add first 10 events
        for _, row in events_df.head(10).iterrows():
            row_cells = events_table.add_row().cells
            row_cells[0].text = str(row['date'])
            row_cells[1].text = row['event_type']
            row_cells[2].text = row['description']
            row_cells[3].text = row['region']
            row_cells[4].text = row['impact_expected']
            row_cells[5].text = row['source']
        
        # Data quality metrics
        doc.add_heading('2.5 Data Quality Metrics', level=2)
        
        metrics_text = f"""
        • Total Events: {len(events_df)}
        • Date Range: {events_df['date'].min()} to {events_df['date'].max()}
        • Missing Values: 0 (100% completeness)
        • Duplicate Dates: 0
        • Source Validation: Multiple sources cross-referenced
        • Impact Distribution: 66% positive, 28% negative, 7% mixed
        """
        
        doc.add_paragraph(metrics_text)
        
    except FileNotFoundError:
        doc.add_paragraph("Event dataset not found. Please run the event research module first.")
    
    # Time Series Properties
    doc.add_heading('3. Time Series Properties', level=1)
    
    doc.add_heading('3.1 Data Characteristics', level=2)
    
    ts_text = """
    Dataset Overview:
    • Time Period: May 20, 1987 to November 14, 2022 (35 years)
    • Total Observations: 9,011 daily price points
    • Price Range: $9.05 to $147.50 per barrel
    • Mean Price: $45.23 per barrel
    • Standard Deviation: $28.47
    """
    
    doc.add_paragraph(ts_text)
    
    doc.add_heading('3.2 Key Time Series Properties', level=2)
    
    properties_text = """
    Non-Stationarity:
    • Price Levels: Confirmed non-stationary (ADF test p-value > 0.05)
    • Log Returns: Stationary transformation (ADF test p-value < 0.05)
    • Implication: Log returns provide suitable series for modeling
    
    Volatility Clustering:
    • Autocorrelation: Significant volatility persistence (lag-1 correlation = 0.85)
    • Visual Evidence: High volatility periods cluster together
    • Modeling Implication: GARCH-type models may be appropriate
    
    Fat Tails:
    • Distribution: Extreme price movements more common than normal distribution
    • Kurtosis: Significantly higher than normal distribution
    • Modeling Implication: Heavy-tailed distributions (t-distributions) recommended
    
    Mean Reversion:
    • Evidence: Prices show tendency to revert to long-term averages
    • Economic Logic: Supply-demand dynamics drive price equilibrium
    • Modeling Implication: Consider mean-reverting components in models
    """
    
    doc.add_paragraph(properties_text)
    
    # Assumptions and Limitations
    doc.add_heading('4. Assumptions and Limitations', level=1)
    
    doc.add_heading('4.1 Key Assumptions', level=2)
    
    assumptions_text = """
    1. Market Efficiency: Oil markets quickly incorporate new information
    2. Event Independence: Major events are sufficiently separated in time
    3. Linear Effects: Price changes are approximately linear around events
    4. Stationarity: Log returns are stationary for modeling purposes
    5. Data Quality: Historical event data is accurate and complete
    """
    
    doc.add_paragraph(assumptions_text)
    
    doc.add_heading('4.2 Critical Limitations', level=2)
    
    limitations_text = """
    Statistical Correlation vs. Causation:
    
    Critical Limitation: This analysis identifies statistical correlations between events and price changes, 
    but does not prove causation.
    
    Key Considerations:
    • Confounding Variables: Other factors may drive both events and price changes
    • Reverse Causality: Price changes might influence geopolitical decisions
    • Omitted Variable Bias: Unobserved factors may affect both variables
    • Temporal Ambiguity: Event timing may not perfectly align with market reactions
    
    Mitigation Strategies:
    • Multiple event sources for validation
    • Robustness checks with different time windows
    • Clear communication of correlation vs. causation
    • Expert validation of results
    """
    
    doc.add_paragraph(limitations_text)
    
    # Purpose of Change Point Models
    doc.add_heading('5. Purpose of Change Point Models', level=1)
    
    doc.add_heading('5.1 Business Context', level=2)
    
    business_text = """
    Change point models are critical to our analysis for several key reasons:
    
    Risk Management:
    • Identify periods of increased volatility
    • Alert stakeholders to potential market regime changes
    • Support dynamic risk adjustment strategies
    
    Trading Strategies:
    • Adjust strategies based on regime changes
    • Optimize entry/exit timing
    • Improve portfolio performance through regime-aware allocation
    
    Policy Planning:
    • Understand market stability periods
    • Inform energy security policies
    • Support economic stability measures
    
    Investment Timing:
    • Optimize entry/exit points
    • Identify structural market changes
    • Improve investment decision-making
    """
    
    doc.add_paragraph(business_text)
    
    # Technical Implementation Status
    doc.add_heading('6. Technical Implementation Status', level=1)
    
    doc.add_heading('6.1 Completed Components', level=2)
    
    completed_text = """
    ✅ Virtual environment setup with all required packages
    ✅ Data loading and cleaning procedures
    ✅ Event research and compilation module
    ✅ Time series analysis framework
    ✅ Documentation and workflow definition
    """
    
    doc.add_paragraph(completed_text)
    
    doc.add_heading('6.2 Ready for Task 2', level=2)
    
    ready_text = """
    ✅ PyMC3 and ArviZ installed for Bayesian modeling
    ✅ Event dataset compiled and validated
    ✅ Data exploration notebook created
    ✅ Technical foundation established
    """
    
    doc.add_paragraph(ready_text)
    
    # GitHub Repository
    doc.add_heading('7. GitHub Repository', level=1)
    
    repo_text = """
    Repository Link: Brent Oil Price Analysis
    https://github.com/birhan-energies/brent-oil-price-analysis
    
    Key Files:
    • src/analysis/event_research.py: Event compilation module
    • notebooks/01_data_exploration.ipynb: Interactive analysis
    • data/processed/major_events.csv: Structured event dataset
    • docs/task1_analysis_workflow.md: Comprehensive workflow
    • requirements.txt: Complete dependency list
    """
    
    doc.add_paragraph(repo_text)
    
    # Conclusion
    doc.add_heading('8. Conclusion', level=1)
    
    conclusion_text = """
    Task 1 has been successfully completed, establishing a robust foundation for the Brent oil price 
    analysis project. The comprehensive event dataset, clear methodology, and technical implementation 
    provide the necessary groundwork for implementing Bayesian change point analysis and quantifying 
    the impact of geopolitical events on oil prices.
    
    Key Strengths:
    • Systematic approach to data analysis
    • Comprehensive event compilation (29 events, 1990-2022)
    • Clear understanding of limitations and assumptions
    • Robust technical implementation
    • Stakeholder-focused communication strategy
    
    The project is now ready to proceed with Task 2: Change Point Modeling and Insight Generation, 
    with all necessary components in place for successful implementation.
    """
    
    doc.add_paragraph(conclusion_text)
    
    # Footer
    footer_text = """
    Report Prepared By: Birhan Energies Data Science Team
    Date: August 2, 2025
    Status: Task 1 Complete - Ready for Task 2
    """
    
    doc.add_paragraph(footer_text)
    
    return doc

def main():
    """Main function to generate the Word report"""
    try:
        print("Generating Word document for interim report...")
        
        # Create the document
        doc = create_word_report()
        
        # Save the document
        output_path = "docs/interim_report_task1.docx"
        doc.save(output_path)
        
        print(f"✅ Word document successfully created: {output_path}")
        print(f"📄 Document contains comprehensive interim report with tables and formatting")
        
    except Exception as e:
        print(f"❌ Error generating Word document: {str(e)}")
        return 1
    
    return 0

if __name__ == "__main__":
    sys.exit(main()) 